"""
Generate Data Coverage & Capability Document (Word .docx)
Japan Culture MCP Server

This script queries the DB and produces an up-to-date .docx report.
Run whenever data is updated to keep the document current.

Usage:
    python3 docs/generate_data_coverage_doc.py
"""
import sqlite3
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# --- Config ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.join(SCRIPT_DIR, "..")
DB_PATH = os.path.join(PROJECT_ROOT, "ontology", "culture_ontology.db")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "Japan_Culture_MCP_Data_Coverage.docx")

# Check for /tmp copy first (faster)
TMP_DB = "/tmp/culture_ontology_p16_final.db"
if os.path.exists(TMP_DB) and os.path.getsize(TMP_DB) > 1_000_000_000:
    DB_PATH = TMP_DB


def open_db():
    db = sqlite3.connect(DB_PATH, timeout=30)
    db.execute("PRAGMA journal_mode=WAL")
    db.execute("PRAGMA busy_timeout=30000")
    db.execute("PRAGMA cache_size=-64000")
    db.execute("PRAGMA mmap_size=268435456")
    return db


def q(db, sql, params=()):
    return db.execute(sql, params).fetchall()


def q1(db, sql, params=()):
    return db.execute(sql, params).fetchone()[0]


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="2B579A"):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def fmt(n):
    """Format number with commas."""
    return f"{n:,}"


def pct(n, total):
    if total == 0:
        return "0.0%"
    return f"{100 * n / total:.1f}%"


def main():
    print(f"Using DB: {DB_PATH}", flush=True)
    db = open_db()
    now = datetime.now()

    # === Gather all stats ===
    print("Gathering stats...", flush=True)

    total_entities = q1(db, "SELECT COUNT(*) FROM entities")
    active = q1(db, "SELECT COUNT(*) FROM entities WHERE is_dormant=0")
    dormant = q1(db, "SELECT COUNT(*) FROM entities WHERE is_dormant=1")
    conns = q1(db, "SELECT COUNT(*) FROM connections")
    geo = q1(db, "SELECT COUNT(*) FROM entities WHERE lat IS NOT NULL AND is_dormant=0")
    total_tags = q1(db, "SELECT COUNT(*) FROM entity_tags")
    tagged_ents = q1(db, "SELECT COUNT(DISTINCT entity_id) FROM entity_tags")
    unique_sources = q1(db, "SELECT COUNT(DISTINCT source) FROM entities WHERE is_dormant=0")
    pilgrim_conns = q1(db, "SELECT COUNT(*) FROM connections WHERE connection_type LIKE 'pilgrimage%'")
    fts_count = q1(db, "SELECT COUNT(*) FROM entities_fts")
    rtree_count = q1(db, "SELECT COUNT(*) FROM entities_rtree")

    # Entity types
    entity_types = q(db, """
        SELECT entity_type, COUNT(*) as cnt,
               SUM(CASE WHEN lat IS NOT NULL THEN 1 ELSE 0 END) as geo_cnt
        FROM entities WHERE is_dormant=0
        GROUP BY entity_type ORDER BY cnt DESC
    """)

    # Source categories
    source_cats_sql = {
        "JapanSearch (SPARQL)": "source LIKE 'jps%' OR source LIKE 'japansearch%'",
        "Wikidata (SPARQL)": "source LIKE 'wd_%' OR source LIKE 'wikidata%'",
        "MADB (SPARQL)": "source LIKE 'madb%'",
        "OpenStreetMap": "source LIKE 'osm%'",
        "国土数値情報": "source LIKE 'kokudo%'",
        "ToMuCo (OAI-PMH)": "source LIKE 'tomuco%'",
        "DBpedia Japanese": "source LIKE 'dbpedia%'",
        "青空文庫": "source LIKE 'aozora%'",
        "ColBase (国立博物館)": "source LIKE 'colbase%'",
        "NDL (国立国会図書館)": "source LIKE 'ndl%'",
    }
    source_cat_data = []
    for cat, where in source_cats_sql.items():
        cnt = q1(db, f"SELECT COUNT(*) FROM entities WHERE is_dormant=0 AND ({where})")
        geo_cnt = q1(db, f"SELECT COUNT(*) FROM entities WHERE is_dormant=0 AND lat IS NOT NULL AND ({where})")
        source_cat_data.append((cat, cnt, geo_cnt))

    # Connection types (top 15)
    conn_types = q(db, """
        SELECT connection_type, COUNT(*) FROM connections
        GROUP BY connection_type ORDER BY COUNT(*) DESC LIMIT 15
    """)

    # Tag axes
    tag_axes = {}
    for axis in ["theme", "era", "medium", "geography", "experience"]:
        tag_axes[axis] = q(db, """
            SELECT value_code, COUNT(*) FROM entity_tags
            WHERE axis=? GROUP BY value_code ORDER BY COUNT(*) DESC
        """, (axis,))

    # External IDs
    ext_ids = {
        "wikidata_id": q1(db, "SELECT COUNT(*) FROM entities WHERE wikidata_id IS NOT NULL AND is_dormant=0"),
        "madb_id": q1(db, "SELECT COUNT(*) FROM entities WHERE madb_id IS NOT NULL AND is_dormant=0"),
        "ndl_id": q1(db, "SELECT COUNT(*) FROM entities WHERE ndl_id IS NOT NULL AND is_dormant=0"),
        "image_url": q1(db, "SELECT COUNT(*) FROM entities WHERE image_url IS NOT NULL AND is_dormant=0"),
    }

    # Pilgrimage breakdown
    pilgrim_types = q(db, """
        SELECT connection_type, COUNT(*) FROM connections
        WHERE connection_type LIKE 'pilgrimage%'
        GROUP BY connection_type ORDER BY COUNT(*) DESC
    """)

    # 5 region stats
    regions = {
        "瀬戸内 (Setouchi)": (33.8, 34.8, 131.5, 135.0),
        "紀伊 (Kii)": (33.5, 34.5, 135.0, 136.5),
        "新潟 (Niigata)": (37.0, 38.5, 138.0, 140.0),
        "京都 (Kyoto)": (34.8, 35.3, 135.3, 136.0),
        "東京 (Tokyo)": (35.5, 35.9, 139.4, 139.9),
    }
    region_stats = []
    for name, (lat_min, lat_max, lon_min, lon_max) in regions.items():
        cnt = q1(db, "SELECT COUNT(*) FROM entities WHERE is_dormant=0 AND lat BETWEEN ? AND ? AND lon BETWEEN ? AND ?",
                 (lat_min, lat_max, lon_min, lon_max))
        region_stats.append((name, cnt))

    db.close()
    print("Stats gathered.", flush=True)

    # === Build Document ===
    print("Building document...", flush=True)
    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Yu Gothic"
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    # === Title Page ===
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Japan Culture MCP Server")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("データカバレッジ & 機能一覧")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"更新日: {now.strftime('%Y-%m-%d')}\nVersion: v1.2.0")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # === 1. Executive Summary ===
    doc.add_heading("1. エグゼクティブサマリー", level=1)
    doc.add_paragraph(
        "Japan Culture MCP Serverは、日本文化に関する包括的なオントロジーデータベースを"
        "Model Context Protocol (MCP) を通じてAIアシスタントに提供するサーバーです。"
    )

    summary_data = [
        ("総エンティティ数", fmt(total_entities)),
        ("アクティブ・エンティティ", fmt(active)),
        ("休眠 (dormant)", fmt(dormant)),
        ("接続 (connections)", fmt(conns)),
        ("座標付きエンティティ", fmt(geo)),
        ("オントロジータグ", fmt(total_tags)),
        ("タグ付きエンティティ", fmt(tagged_ents)),
        ("データソース数", fmt(unique_sources)),
        ("聖地巡礼接続", fmt(pilgrim_conns)),
        ("MCPツール数", "39"),
        ("FTS5全文検索", fmt(fts_count) + " entries"),
        ("R-Tree空間索引", fmt(rtree_count) + " entries"),
    ]
    add_styled_table(doc, ["指標", "値"], summary_data, col_widths=[6, 6])

    # === 2. Data Sources ===
    doc.add_page_break()
    doc.add_heading("2. データソース一覧とカバレッジ", level=1)
    doc.add_paragraph(
        "以下は主要データソース別のエンティティ数と座標付与率です。"
        "ソースは取得APIまたは元データベースで分類しています。"
    )

    source_rows = []
    for cat, cnt, geo_cnt in source_cat_data:
        source_rows.append((
            cat, fmt(cnt), fmt(geo_cnt),
            pct(geo_cnt, cnt), pct(cnt, active)
        ))
    add_styled_table(
        doc,
        ["データソース", "エンティティ数", "座標付き", "座標率", "全体比"],
        source_rows,
        col_widths=[5.5, 3, 2.5, 2, 2],
    )

    doc.add_paragraph()
    doc.add_heading("2.1 データソース詳細", level=2)

    source_details = [
        ("JapanSearch", "ジャパンサーチ SPARQL API",
         "版画、書籍、古文書、写真、新聞、音楽、映像など264以上の文化機関DBを横断。"
         "rdfs:labelページネーション + schema:datePublished日付範囲クエリで大量取得。"),
        ("Wikidata", "Wikidata SPARQL",
         "神社仏閣、スポーツ選手、音楽、人物、企業、映画、ゲーム、キャラクター。"
         "P840/P915で聖地巡礼データも取得。10s wait + 504 retry。"),
        ("MADB", "メディア芸術データベース SPARQL",
         "漫画（25万冊）、アニメ（9千タイトル）、ゲーム（3.5万）の構造化メタデータ。"),
        ("OpenStreetMap", "Overpass API",
         "寺社、鳥居、文化的ランドマーク。地域分割でタイムアウト回避。全件座標付き。"),
        ("国土数値情報", "GeoJSON/Shapefile",
         "観光スポット、文化財、世界遺産、観光施設。全件座標付き。"),
        ("ToMuCo", "OAI-PMH",
         "東京の博物館コレクション（東京国立博物館等）。約86%が座標付き。"),
        ("DBpedia Japanese", "SPARQL",
         "場所、人物、イベント、作品の属性情報。"),
        ("青空文庫", "Web Scraping",
         "古典・近代日本文学 約16,000作品。著作権切れテキスト。"),
        ("ColBase", "API",
         "国立博物館コレクション（東博、京博、奈良博、九博）。"),
        ("NDL", "SRU/IIIF",
         "国立国会図書館の古典籍、浮世絵、デジタルコレクション。"),
    ]

    for name, api, desc in source_details:
        p = doc.add_paragraph()
        run = p.add_run(f"{name} ({api}): ")
        run.bold = True
        run.font.size = Pt(9)
        run2 = p.add_run(desc)
        run2.font.size = Pt(9)

    # === 3. Entity Types ===
    doc.add_page_break()
    doc.add_heading("3. エンティティ種別とカバレッジ", level=1)
    doc.add_paragraph(
        "各エンティティ種別の件数、座標付与率、全体に占める割合です。"
        "「work」が全体の82%を占め、JapanSearchからの書籍・版画・新聞が中心です。"
    )

    # Only show meaningful types (>100)
    et_rows = []
    for etype, cnt, geo_cnt in entity_types:
        if cnt < 100:
            continue
        et_rows.append((
            str(etype), fmt(cnt), fmt(geo_cnt),
            pct(geo_cnt, cnt), pct(cnt, active)
        ))
    add_styled_table(
        doc,
        ["エンティティ種別", "件数", "座標付き", "座標率", "全体比"],
        et_rows,
        col_widths=[4.5, 3, 2.5, 2, 2],
    )

    # === 4. Ontology (5-axis tags) ===
    doc.add_page_break()
    doc.add_heading("4. 5軸オントロジー", level=1)
    doc.add_paragraph(
        f"エンティティには5軸のタグが付与されています（全{fmt(total_tags)}タグ、"
        f"{fmt(tagged_ents)}エンティティ）。各軸の値と件数は以下の通りです。"
    )

    axis_ja = {
        "theme": ("テーマ", "文化的テーマ・モチーフ"),
        "era": ("時代", "歴史的時代区分"),
        "medium": ("媒体", "表現媒体・ジャンル"),
        "geography": ("地理", "地域分類"),
        "experience": ("体験", "体験モード"),
    }

    for axis, values in tag_axes.items():
        ja_name, ja_desc = axis_ja[axis]
        doc.add_heading(f"4.{list(tag_axes.keys()).index(axis)+1} {ja_name} ({axis}) — {ja_desc}", level=2)
        ax_rows = [(str(vc), fmt(cnt)) for vc, cnt in values]
        add_styled_table(doc, ["値コード", "タグ件数"], ax_rows, col_widths=[6, 4])
        doc.add_paragraph()

    # === 5. Connections ===
    doc.add_page_break()
    doc.add_heading("5. 接続 (Connections)", level=1)
    doc.add_paragraph(
        f"エンティティ間の文化的接続は全{fmt(conns)}件です。"
        "接続は5軸距離スコア（theme/era/medium/geography/experience）と"
        "セレンディピティスコアを持ちます。"
    )

    ct_rows = [(str(ct), fmt(cnt)) for ct, cnt in conn_types]
    add_styled_table(doc, ["接続タイプ", "件数"], ct_rows, col_widths=[8, 4])

    doc.add_heading("5.1 聖地巡礼接続", level=2)
    doc.add_paragraph(
        f"聖地巡礼関連の接続は全{fmt(pilgrim_conns)}件で、"
        "アニメ・映画・ゲーム等の作品と実在の場所を結びます。"
    )
    pt_rows = [(str(ct), fmt(cnt)) for ct, cnt in pilgrim_types]
    add_styled_table(doc, ["接続タイプ", "件数"], pt_rows, col_widths=[8, 4])

    # === 6. Geo Coverage ===
    doc.add_page_break()
    doc.add_heading("6. 地理的カバレッジ", level=1)
    doc.add_paragraph(
        f"座標付きエンティティは{fmt(geo)}件（アクティブの{pct(geo, active)}）です。"
        "R-Tree空間インデックスによる高速な近傍検索が可能です。"
    )

    doc.add_heading("6.1 主要地域のサンプル数", level=2)
    reg_rows = [(name, fmt(cnt)) for name, cnt in region_stats]
    add_styled_table(doc, ["地域", "エンティティ数"], reg_rows, col_widths=[6, 4])

    # === 7. External IDs ===
    doc.add_heading("6.2 外部ID紐付け", level=2)
    id_rows = [(k, fmt(v), pct(v, active)) for k, v in ext_ids.items()]
    add_styled_table(doc, ["外部ID", "紐付き件数", "カバレッジ"], id_rows, col_widths=[4, 4, 3])

    # === 8. Tools ===
    doc.add_page_break()
    doc.add_heading("7. MCPツール一覧 (39ツール)", level=1)
    doc.add_paragraph(
        "以下はMCPプロトコル経由で利用可能な全ツールです。"
    )

    tools = [
        ("コアツール", [
            ("search_anime", "AniList GraphQL APIでアニメ・漫画検索"),
            ("search_media_arts", "MADB SPARQLで漫画・アニメ・ゲーム検索"),
            ("cross_reference", "AniList×MADB統合検索"),
            ("search_japan_search", "ジャパンサーチ横断検索（264機関）"),
            ("search_wikidata", "Wikidata日本文化エンティティ検索"),
            ("resolve_entity", "名前→Wikidata ID解決"),
            ("get_ndl_manifest", "NDL IIIFマニフェスト取得"),
            ("get_ndl_ocr_text", "NDL OCRテキスト取得"),
            ("search_ndl", "NDLサーチSRU検索"),
            ("search_dbpedia_ja", "DBpedia Japanese検索"),
            ("iiif_get_manifest", "汎用IIIFマニフェスト取得"),
            ("get_map_tile_url", "国土地理院タイルURL生成"),
            ("get_heritage_map_url", "文化財総覧WebGIS URL生成"),
            ("get_tourism_stats", "e-Stat観光統計取得"),
            ("cross_reference_v2", "全データソース横断検索"),
        ]),
        ("セレンディピティ・発見", [
            ("find_serendipity", "文化的セレンディピティ発見"),
            ("explore_axis", "5軸オントロジー探索"),
            ("get_entity_detail", "エンティティ詳細取得"),
            ("get_cultural_route", "文化ルート生成"),
            ("search_culture", "横断検索シンプル版"),
        ]),
        ("特化型検索", [
            ("search_traditional_crafts", "伝統工芸検索"),
            ("search_literature", "青空文庫文学検索"),
            ("search_artworks", "美術作品横断検索"),
            ("search_festivals", "祭り・無形文化遺産検索"),
            ("search_living_national_treasures", "人間国宝検索"),
            ("generate_serendipity_route", "セレンディピティルート生成"),
            ("explore_connections", "接続グラフBFS探索"),
            ("get_culture_stats", "DB統計取得"),
        ]),
        ("聖地巡礼・位置情報", [
            ("search_pilgrimage", "聖地巡礼スポット検索"),
            ("generate_pilgrimage_route", "聖地巡礼ルート生成"),
            ("get_nearby_culture", "座標周辺文化リソース検索"),
        ]),
        ("分析・比較", [
            ("generate_timeline", "文化タイムライン生成"),
            ("compare_cultures", "2文化要素比較"),
            ("generate_culture_map", "GeoJSON文化地図生成"),
            ("today_in_culture", "今日の文化トピック"),
            ("deep_dive", "エンティティ深掘り推薦"),
        ]),
        ("観光分析", [
            ("get_region_profile", "地域文化プロファイル生成"),
            ("find_tourism_assets", "観光文化資産一覧"),
            ("analyze_cultural_density", "文化密度ヒートマップ"),
        ]),
    ]

    for category, tool_list in tools:
        doc.add_heading(f"7.x {category}", level=2)
        t_rows = [(name, desc) for name, desc in tool_list]
        add_styled_table(doc, ["ツール名", "説明"], t_rows, col_widths=[5.5, 9])
        doc.add_paragraph()

    # === 9. What you can do ===
    doc.add_page_break()
    doc.add_heading("8. このMCPで何ができるか", level=1)

    capabilities = [
        ("8.1 文化横断検索",
         "264以上の文化機関DBを一括検索。アニメ・漫画から浮世絵・古典籍まで、"
         "ジャンルを超えた日本文化リソースの発見が可能。"
         f"対象: {fmt(active)}エンティティ、{fmt(unique_sources)}データソース。"),
        ("8.2 セレンディピティ発見",
         "5軸オントロジー（テーマ・時代・媒体・地理・体験）に基づく"
         "「意外な文化的つながり」の発見。例: 北斎→ポニョ（波のモチーフ共有）。"
         f"接続グラフ: {fmt(conns)}エッジ。"),
        ("8.3 聖地巡礼",
         "アニメ・映画の聖地巡礼スポット検索とルート生成。"
         "近隣の伝統文化スポット（寺社・文化財）も自動推薦。"
         f"聖地接続: {fmt(pilgrim_conns)}件。Google Maps連携対応。"),
        ("8.4 地理空間分析",
         f"座標付き{fmt(geo)}エンティティのR-Tree空間インデックスによる"
         "高速な近傍検索。文化密度ヒートマップの生成、"
         "地域プロファイル（エンティティ統計・テーマ分布・接続密度）の分析。"),
        ("8.5 文化比較・タイムライン",
         "2つの文化要素の共通点・相違点・意外な接続の分析。"
         "テーマ別の文化的時系列生成（時代/地域フィルタ対応）。"),
        ("8.6 観光分析",
         "地域の文化プロファイル生成、観光資産のカテゴリ別一覧、"
         "格子状の文化密度分析（ヒートマップ可視化用GeoJSON出力）。"),
        ("8.7 全文検索 (FTS5)",
         f"日本語・英語テキストの高速全文検索（{fmt(fts_count)}エンティティ対象）。"
         "LIKE検索の225倍高速（4ms vs 900ms）。"),
        ("8.8 伝統文化深掘り",
         "伝統工芸（経産省指定244品目+関連）、人間国宝、祭り・無形文化遺産、"
         "美術作品（ToMuCo 35K + ColBase + 国宝・重文16K）の専門検索。"),
    ]

    for title, desc in capabilities:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    # === 10. Known Limitations ===
    doc.add_page_break()
    doc.add_heading("9. 既知の制約・今後の拡張予定", level=1)

    limitations = [
        ("時間データの不足",
         "entities テーブルに release_year カラムがない（Phase 17で追加予定）。"
         "時系列分析には era タグ（10時代区分）のみ利用可能。"),
        ("anilist_id 未紐付け",
         "AniList GraphQL APIとの直接JOIN不可（全件NULL）。"
         "Phase 17 Step 3 で AniList JSON マッチングにより紐付け予定。"),
        ("ポップカルチャーの座標不足",
         "anime/manga タグ付きで座標付きエンティティは約850件と少ない。"
         "聖地巡礼接続経由での間接的な地理分析が現実的。"),
        ("文化財指定ランク未分類",
         "国宝・重要文化財・県指定の区分フィールドなし。"
         "search_artworks ツール経由で16K件にアクセスは可能。"),
        ("地理粒度",
         "geography タグは13地域ブロック（東京/関東/中部/近畿...）。"
         "都道府県レベルの分析には座標ベースの自前計算が必要。"),
    ]

    for title, desc in limitations:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # === 11. Update Log ===
    doc.add_heading("10. 更新履歴", level=1)
    doc.add_paragraph(
        f"{now.strftime('%Y-%m-%d')} — 初版作成 "
        f"(v1.2.0, {fmt(total_entities)}エンティティ, {fmt(conns)}接続)"
    )
    doc.add_paragraph(
        "※ このドキュメントは docs/generate_data_coverage_doc.py を実行して自動生成されます。"
        "データ増加時は必ず再実行してください。"
    )

    # === Save ===
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved: {OUTPUT_PATH}", flush=True)
    print(f"  Pages: ~10", flush=True)
    print(f"  Data: {fmt(total_entities)} entities, {fmt(conns)} connections", flush=True)
    print("Done.", flush=True)


if __name__ == "__main__":
    main()
